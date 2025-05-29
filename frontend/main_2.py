# streamlit_app/main.py
import streamlit as st
import requests
import os
import io
import docx
import pandas as pd
import PyPDF2

API_BASE = os.getenv("API_BASE", "http://localhost:8000")

# Session state keys
if "document_id" not in st.session_state:
    st.session_state.document_id = None
if "paragraphs" not in st.session_state:
    st.session_state.paragraphs = []
if "rules" not in st.session_state:
    st.session_state.rules = []
if "violations" not in st.session_state:
    st.session_state.violations = []
if "current_violation_index" not in st.session_state:
    st.session_state.current_violation_index = 0
if "highlighted_text" not in st.session_state:
    st.session_state.highlighted_text = ""
if "suggested_fix" not in st.session_state:
    st.session_state.suggested_fix = ""

st.set_page_config(layout="wide")
st.title("🛡️ Document Compliance Checker")

# --- LEFT SIDEBAR ---

with st.sidebar:
    st.header("📜 Compliance Rules")

    # Fetch compliance rules
    if not st.session_state.rules:
        r = requests.get(f"{API_BASE}/rules")  # <- You should expose this endpoint
        if r.ok:
            st.session_state.rules = r.json()

    selected_rule_name = None
    selected_rule_id = None

    for rule in st.session_state.rules:
        if st.button(rule['name']):
            selected_rule_name = rule['name']
            selected_rule_id = rule['id']
            st.session_state.current_violation_index = 0
            st.session_state.violations = []
            st.session_state.suggested_fix = ""
            st.session_state.highlighted_text = ""

    if selected_rule_id:
        st.subheader(f"✏️ Violations for: {selected_rule_name}")

        # Fetch violations only once
        if not st.session_state.violations:
            for para in st.session_state.paragraphs:
                res = requests.post(f"{API_BASE}/check_violation", json={
                    "rule_id": selected_rule_id,
                    "paragraph_id": para["id"]
                })
                if res.ok:
                    result = res.json()
                    st.session_state.violations.append({
                        "violation_id": result["violation_id"],
                        "paragraph_id": para["id"],
                        "highlighted": result["highlighted_text"]
                    })

        if st.session_state.violations:
            curr = st.session_state.violations[st.session_state.current_violation_index]
            st.session_state.highlighted_text = curr["highlighted"]

            # LLM fix suggestion
            if not st.session_state.suggested_fix:
                fix = requests.post(f"{API_BASE}/suggest_fix", params={
                    "violation_id": curr["violation_id"]
                })
                if fix.ok:
                    st.session_state.suggested_fix = fix.json()["suggested_fix"]

            st.text_area("✂️ LLM Suggested Fix:", value=st.session_state.suggested_fix, key="manual_fix", height=150)

            if st.button("✅ Apply Fix"):
                # Accept edit and update backend
                new_text = st.session_state.manual_fix
                violation_id = curr["violation_id"]
                resp = requests.post(f"{API_BASE}/accept_edit", json={
                    "violation_id": violation_id,
                    "new_text": new_text,
                    "accepted": True
                })
                if resp.ok:
                    # Reload paragraphs from server
                    r = requests.get(f"{API_BASE}/document_paragraphs", params={"doc_id": st.session_state.document_id})
                    if r.ok:
                        st.session_state.paragraphs = r.json()

                    st.session_state.current_violation_index += 1
                    st.session_state.suggested_fix = ""
                    st.session_state.highlighted_text = ""

                    if st.session_state.current_violation_index >= len(st.session_state.violations):
                        st.success("✅ All violations processed!")

# --- MAIN CONTENT AREA ---

def extract_text(file):
    file_type = uploaded.name.split('.')[-1].lower()

    if file_type == "txt":
        return file.read().decode("utf-8")

    elif file_type == "pdf":
        reader = PyPDF2.PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif file_type == "csv":
        df = pd.read_csv(file)
        return df.to_string(index=False)

    elif file_type == "docx":
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        st.error("Unsupported file format.")
        return ""

col1, col2 = st.columns([1, 3])

with col1:
    st.subheader("📁 Upload Document")
    uploaded = st.file_uploader( "Upload a document", type=["txt", "pdf", "csv", "docx"])

    # Upload & parse
    if uploaded and st.button("Upload & Parse"):
        content = extract_text(uploaded)

        if content:
            res = requests.post(f"{API_BASE}/upload", json={"name": uploaded.name, "content": content})
            if res.ok:
                st.session_state.document_id = res.json()["document_id"]
                st.success("Document uploaded!")

                # Fetch paragraphs
                r = requests.get(f"{API_BASE}/document_paragraphs", params={"doc_id": st.session_state.document_id})
                if r.ok:
                    st.session_state.paragraphs = r.json()
                    st.session_state.violations = []
                    st.session_state.current_violation_index = 0


with col2:
    st.subheader("📄 Document Viewer")

    for para in st.session_state.paragraphs:
        para_text = para["content"]
        # Highlight if it matches current highlighted
        if para["id"] == st.session_state.violations[st.session_state.current_violation_index]["paragraph_id"] if st.session_state.violations else False:
            para_text = para_text.replace(
                st.session_state.highlighted_text,
                f"**:red[{st.session_state.highlighted_text}]**"
            )
        st.markdown(para_text, unsafe_allow_html=True)
